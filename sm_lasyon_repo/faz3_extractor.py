#!/usr/bin/env python3
"""
FAZ-3: PDF, DOCX, Resim Dosyalarını Okuyucu ve Analiz Edici
Tüm PDF/DOCX/JPG/PNG/JFIF dosyalarını tarar, içerik çıkarır ve JSON'a kaydeder.
"""
import os
import json
import sys
import traceback

REPO_DIR = r"C:\Users\soldi\AppData\Local\Temp\Untitled65.ipynb4156691301457397376\sm_lasyon_repo"
OUTPUT_JSON = os.path.join(REPO_DIR, "faz3_extracted_data.json")

results = {
    "pdf_files": {},
    "docx_files": {},
    "image_files": {},
    "summary": {
        "total_pdfs": 0,
        "total_docxs": 0,
        "total_images": 0,
        "pdfs_read": 0,
        "pdfs_failed": 0,
        "docxs_read": 0,
        "docxs_failed": 0,
        "images_read": 0,
        "images_failed": 0
    }
}

# ============================================================
# PDF OKUMA
# ============================================================
def read_pdf_pypdf2(filepath):
    """PyPDF2 ile PDF oku"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        num_pages = len(reader.pages)
        text_parts = []
        for i, page in enumerate(reader.pages):
            try:
                t = page.extract_text()
                if t:
                    text_parts.append(f"--- Sayfa {i+1} ---\n{t}")
            except:
                text_parts.append(f"--- Sayfa {i+1} ---\n[OKUNAMADI]")
        return num_pages, "\n".join(text_parts)
    except Exception as e:
        return -1, f"PyPDF2 HATA: {str(e)}"

def read_pdf_pdfplumber(filepath):
    """pdfplumber ile PDF oku (yedek)"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            num_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                try:
                    t = page.extract_text()
                    if t:
                        text_parts.append(f"--- Sayfa {i+1} ---\n{t}")
                except:
                    text_parts.append(f"--- Sayfa {i+1} ---\n[OKUNAMADI]")
        return num_pages, "\n".join(text_parts)
    except Exception as e:
        return -1, f"pdfplumber HATA: {str(e)}"

def extract_numerical_values(text):
    """Metinden sayısal değerleri, formülleri, frekansları çıkar"""
    import re
    findings = {
        "frequencies_hz": [],
        "coordinates": [],
        "constants": [],
        "formulas": [],
        "numbers_11": [],
        "pi_values": [],
        "lambda_values": [],
        "omega_values": []
    }
    
    # Frekans değerleri (Hz, MHz, GHz)
    freq_patterns = [
        r'(\d+\.?\d*)\s*(?:Hz|MHz|GHz|kHz|THz)',
        r'(?:frekans|frequency|rezonans|resonance)[:\s]*(\d+\.?\d*)',
    ]
    for pat in freq_patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            findings["frequencies_hz"].append(m.group(0))
    
    # Koordinatlar
    coord_patterns = [
        r'(\d+°\d+\'\d+\.?\d*"[NS]?\s*,?\s*\d+°\d+\'\d+\.?\d*"[EW]?)',
        r'(?:koordinat|coordinate|enlem|boylam)[:\s]*(\d+\.?\d*[°]?\s*[NS]?\s*,?\s*\d+\.?\d*[°]?\s*[EW]?)',
    ]
    for pat in coord_patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            findings["coordinates"].append(m.group(0))
    
    # Sabitler
    const_patterns = [
        r'(?:sabit|constant|değer|value)[:\s]*(\d+\.?\d*(?:e[+-]?\d+)?)',
        r'(?:G\s*=\s*|c\s*=\s*|h\s*=\s*|ħ\s*=\s*|kB\s*=\s*)(\d+\.?\d*(?:e[+-]?\d+)?)',
    ]
    for pat in const_patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            findings["constants"].append(m.group(0))
    
    # Formüller
    formula_patterns = [
        r'[A-Za-z]+\s*=\s*[A-Za-z0-9\s\+\-\*/\^\(\)\.π√Σ∫∂]+',
        r'(?:formül|formula|denklem|equation)[:\s]*([^\n]+)',
    ]
    for pat in formula_patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            findings["formulas"].append(m.group(0)[:200])
    
    # 11 ile ilgili sayılar
    for m in re.finditer(r'\b(?:11|111|1111|11111|111111|1111111|11111111|111111111)\b', text):
        findings["numbers_11"].append(m.group(0))
    
    # Pi değerleri
    for m in re.finditer(r'(?:pi|π|Pi_11|PI)[:\s]*(\d+\.?\d*)', text, re.IGNORECASE):
        findings["pi_values"].append(m.group(0))
    
    # Lambda değerleri
    for m in re.finditer(r'(?:lambda|λ|Lambda)[:\s]*(\d+\.?\d*)', text, re.IGNORECASE):
        findings["lambda_values"].append(m.group(0))
    
    # Omega değerleri
    for m in re.finditer(r'(?:omega|Ω|Omega)[:\s]*(\d+\.?\d*)', text, re.IGNORECASE):
        findings["omega_values"].append(m.group(0))
    
    # Tekrar edenleri temizle
    for k in findings:
        findings[k] = list(set(findings[k]))[:50]  # max 50
    
    return findings

def analyze_pdf_content(text):
    """PDF içeriğinden anlamlı bilgiler çıkar"""
    analysis = {
        "has_formulas": False,
        "has_coordinates": False,
        "has_frequencies": False,
        "has_11_reference": False,
        "has_pi_reference": False,
        "has_lambda_reference": False,
        "has_omega_reference": False,
        "has_hudhud": False,
        "has_kailash": False,
        "has_giza": False,
        "has_halley": False,
        "has_malta": False,
        "has_levhi_mahfuz": False,
        "has_antigravity": False,
        "has_dark_energy": False,
        "has_simulation": False,
        "has_quantum": False,
        "has_11d": False,
        "has_celali": False,
        "has_orhun": False,
        "has_maya": False,
        "has_elon": False,
        "has_nasa": False,
        "has_consciousness": False,
        "has_dna": False,
        "has_pineal": False,
        "has_ley_line": False,
        "has_geoid": False,
        "has_warp": False,
        "has_wormhole": False,
        "has_entanglement": False,
        "has_casimir": False,
        "has_planck": False,
        "has_string": False,
        "has_multiverse": False,
        "has_graviton": False,
        "has_tachyon": False,
        "has_axion": False,
        "has_monopole": False,
        "has_inflation": False,
        "has_heat_death": False,
        "has_golden_ratio": False,
        "has_euler": False,
        "has_fibonacci": False,
        "has_repunit": False,
        "has_666": False,
        "has_814": False,
        "has_22_66_88": False,
        "has_151": False,
        "has_369": False,
        "has_tesla": False,
        "has_vibration": False,
        "has_akashic": False,
        "has_noosphere": False,
        "has_observer": False,
        "has_bridge": False,
        "has_heart_brain": False,
        "has_spine": False,
        "has_merkaba": False,
        "has_torus": False,
        "has_flower_of_life": False,
        "has_metatron": False,
        "has_sephirot": False,
        "has_emerald_tablet": False,
        "has_hermes": False,
        "has_thoth": False,
        "has_enoch": False,
        "has_solomon": False,
        "has_ark": False,
        "has_flood": False,
        "has_atlantis": False,
        "has_mu": False,
        "has_lemuria": False,
        "has_anunnaki": False,
        "has_nibiru": False,
        "has_sirius": False,
        "has_orion": False,
        "has_pleiades": False,
        "has_andromeda": False,
        "has_milky_way": False,
        "has_sun": False,
        "has_moon": False,
        "has_earth": False,
        "has_mars": False,
        "has_jupiter": False,
        "has_saturn": False,
        "has_venus": False,
        "has_mercury": False,
        "has_uranus": False,
        "has_neptune": False,
        "has_pluto": False,
        "has_planet_x": False,
        "has_asteroid": False,
        "has_comet": False,
        "has_meteor": False,
        "has_eclipse": False,
        "has_solstice": False,
        "has_equinox": False,
        "has_precession": False,
        "has_axial_tilt": False,
        "has_magnetic_field": False,
        "has_schumann": False,
        "has_van_allen": False,
        "has_ionosphere": False,
        "has_magnetosphere": False,
        "has_tectonic": False,
        "has_volcano": False,
        "has_earthquake": False,
        "has_tsunami": False,
        "has_pole_shift": False,
        "has_ice_age": False,
        "has_global_warming": False,
        "has_climate": False,
        "has_carbon": False,
        "has_oxygen": False,
        "has_hydrogen": False,
        "has_helium": False,
        "has_nitrogen": False,
        "has_silicon": False,
        "has_crystal": False,
        "has_quartz": False,
        "has_water": False,
        "has_fire": False,
        "has_air": False,
        "has_earth_element": False,
        "has_ether": False,
        "has_akasha": False,
        "has_prana": False,
        "has_chi": False,
        "has_kundalini": False,
        "has_chakra": False,
        "has_third_eye": False,
        "has_crown": False,
        "has_root": False,
        "has_heart": False,
        "has_throat": False,
        "has_solar_plexus": False,
        "has_sacral": False,
    }
    
    text_lower = text.lower()
    
    keyword_map = {
        "has_formulas": ["formül", "formula", "denklem", "equation", "="],
        "has_coordinates": ["koordinat", "coordinate", "enlem", "boylam", "latitude", "longitude"],
        "has_frequencies": ["frekans", "frequency", "hz", "mhz", "ghz", "rezonans", "resonance"],
        "has_11_reference": [" 11 ", "11.", "11-", "11d", "11 boyut", "11 dimension"],
        "has_pi_reference": ["pi", "π", "3.14"],
        "has_lambda_reference": ["lambda", "λ", "6.666"],
        "has_omega_reference": ["omega", "Ω", "ω"],
        "has_hudhud": ["hüdhüd", "hudhud", "hud hud", "ibibik", "hoopoe"],
        "has_kailash": ["kailaş", "kailash", "kailaşa"],
        "has_giza": ["giza", "gize", "piramit", "pyramid", "keops"],
        "has_halley": ["halley", "halley"],
        "has_malta": ["malta"],
        "has_levhi_mahfuz": ["levh", "mahfuz", "mahfus", "lehf"],
        "has_antigravity": ["anti-yerçekimi", "antigravity", "anti gravity", "anti gravit"],
        "has_dark_energy": ["karanlık enerji", "dark energy", "karanlık madde", "dark matter"],
        "has_simulation": ["simülasyon", "simulation", "simülat"],
        "has_quantum": ["kuantum", "quantum", "dalga", "wave", "parçacık", "particle"],
        "has_11d": ["11d", "11-boyut", "11 dimension", "11 dimensional"],
        "has_celali": ["celali", "celal", "celali"],
        "has_orhun": ["orhun", "orhun", "göktürk", "yazıt"],
        "has_maya": ["maya", "mayan"],
        "has_elon": ["elon", "musk"],
        "has_nasa": ["nasa"],
        "has_consciousness": ["bilinç", "consciousness", "farkındalık", "awareness"],
        "has_dna": ["dna", "rna", "genetik", "genetic"],
        "has_pineal": ["pineal", "epifiz", "üçüncü göz", "third eye"],
        "has_ley_line": ["ley hatt", "ley line", "enerji hatt"],
        "has_geoid": ["geoid", "geo", "matris", "matrix"],
        "has_warp": ["warp", "alcubierre", "ışık hızı", "ftl"],
        "has_wormhole": ["solucan deliği", "wormhole", "einstein-rosen"],
        "has_entanglement": ["dolanıklık", "entanglement", "epr"],
        "has_casimir": ["casimir", "vakum enerjisi", "vacuum energy"],
        "has_planck": ["planck", "planck"],
        "has_string": ["string", "sicim", "süpersicim", "superstring"],
        "has_multiverse": ["çoklu evren", "multiverse", "paralel evren"],
        "has_graviton": ["graviton", "kütleçekim", "gravity"],
        "has_tachyon": ["takyon", "tachyon"],
        "has_axion": ["aksiyon", "axion"],
        "has_monopole": ["monopol", "monopole"],
        "has_inflation": ["enflasyon", "inflation", "şişme"],
        "has_heat_death": ["ısı ölümü", "heat death", "entropi", "entropy"],
        "has_golden_ratio": ["altın oran", "golden ratio", "phi", "φ", "1.618"],
        "has_euler": ["euler", "e sayısı", "2.718"],
        "has_fibonacci": ["fibonacci", "fibonacci"],
        "has_repunit": ["repunit", "1-11-111", "111111"],
        "has_666": ["666", "6.666"],
        "has_814": ["814"],
        "has_22_66_88": ["22-66-88", "22 66 88"],
        "has_151": ["151", "151.993"],
        "has_369": ["369", "3-6-9", "tesla"],
        "has_tesla": ["tesla", "nikola"],
        "has_vibration": ["titreşim", "vibration", "frekans"],
        "has_akashic": ["akaşik", "akashic", "akash"],
        "has_noosphere": ["noosfer", "noosphere"],
        "has_observer": ["gözlemci", "observer", "kuantum gözlem"],
        "has_bridge": ["köprü", "bridge"],
        "has_heart_brain": ["kalp-beyin", "heart-brain", "heartmath"],
        "has_spine": ["omurga", "spine", "omurilik"],
        "has_merkaba": ["merkaba", "merkabah"],
        "has_torus": ["torus", "toroid", "toroidal"],
        "has_flower_of_life": ["yaşam çiçeği", "flower of life"],
        "has_metatron": ["metatron"],
        "has_sephirot": ["sefirot", "sephirot", "hayat ağacı"],
        "has_emerald_tablet": ["zümrüt tablet", "emerald tablet"],
        "has_hermes": ["hermes", "hermetik"],
        "has_thoth": ["thoth", "toth"],
        "has_enoch": ["enoch", "idris", "hanok"],
        "has_solomon": ["süleyman", "solomon"],
        "has_ark": ["nuh", "gemi", "ark", "tufan"],
        "has_flood": ["tufan", "flood", "sel"],
        "has_atlantis": ["atlantis"],
        "has_mu": ["mu kıtası", "mu continent"],
        "has_lemuria": ["lemurya", "lemuria"],
        "has_anunnaki": ["anunnaki", "anunaki"],
        "has_nibiru": ["nibiru", "planet x"],
        "has_sirius": ["sirius", "sirius"],
        "has_orion": ["orion", "avcı"],
        "has_pleiades": ["pleiades", "ülker", "süreyya"],
        "has_andromeda": ["andromeda"],
        "has_milky_way": ["samanyolu", "milky way"],
        "has_sun": ["güneş", "sun", "solar"],
        "has_moon": ["ay", "moon", "lunar"],
        "has_earth": ["dünya", "earth", "yerküre"],
        "has_mars": ["mars", "merih"],
        "has_jupiter": ["jüpiter", "jupiter", "müşteri"],
        "has_saturn": ["satürn", "saturn", "zuhal"],
        "has_venus": ["venüs", "venus", "zühre"],
        "has_mercury": ["merkür", "mercury", "utarit"],
        "has_uranus": ["uranüs", "uranus"],
        "has_neptune": ["neptün", "neptune"],
        "has_pluto": ["plüton", "pluto"],
        "has_planet_x": ["planet x", "gezegen x"],
        "has_asteroid": ["asteroit", "asteroid"],
        "has_comet": ["kuyruklu yıldız", "comet", "komet"],
        "has_meteor": ["meteor", "göktaşı"],
        "has_eclipse": ["tutulma", "eclipse"],
        "has_solstice": ["gündönümü", "solstice"],
        "has_equinox": ["ekinoks", "equinox"],
        "has_precession": ["presesyon", "precession", "yalpalama"],
        "has_axial_tilt": ["eksen eğikliği", "axial tilt"],
        "has_magnetic_field": ["manyetik alan", "magnetic field"],
        "has_schumann": ["schumann", "şuman"],
        "has_van_allen": ["van allen"],
        "has_ionosphere": ["iyonosfer", "ionosphere"],
        "has_magnetosphere": ["manyetosfer", "magnetosphere"],
        "has_tectonic": ["tektonik", "tectonic"],
        "has_volcano": ["volkan", "volcano", "yanardağ"],
        "has_earthquake": ["deprem", "earthquake", "zelzele"],
        "has_tsunami": ["tsunami"],
        "has_pole_shift": ["kutup kayması", "pole shift"],
        "has_ice_age": ["buz çağı", "ice age"],
        "has_global_warming": ["küresel ısınma", "global warming"],
        "has_climate": ["iklim", "climate"],
        "has_carbon": ["karbon", "carbon"],
        "has_oxygen": ["oksijen", "oxygen"],
        "has_hydrogen": ["hidrojen", "hydrogen"],
        "has_helium": ["helyum", "helium"],
        "has_nitrogen": ["azot", "nitrogen"],
        "has_silicon": ["silikon", "silicon", "silisyum"],
        "has_crystal": ["kristal", "crystal"],
        "has_quartz": ["kuvars", "quartz"],
        "has_water": ["su", "water", "h2o"],
        "has_fire": ["ateş", "fire"],
        "has_air": ["hava", "air"],
        "has_earth_element": ["toprak", "earth element"],
        "has_ether": ["eter", "ether", "esir"],
        "has_akasha": ["akaşa", "akasha"],
        "has_prana": ["prana", "yaşam enerjisi"],
        "has_chi": ["chi", "qi", "ki"],
        "has_kundalini": ["kundalini"],
        "has_chakra": ["çakra", "chakra"],
        "has_third_eye": ["üçüncü göz", "third eye", "pineal"],
        "has_crown": ["taç çakra", "crown chakra"],
        "has_root": ["kök çakra", "root chakra"],
        "has_heart": ["kalp çakra", "heart chakra"],
        "has_throat": ["boğaz çakra", "throat chakra"],
        "has_solar_plexus": ["solar pleksus", "solar plexus"],
        "has_sacral": ["sakral", "sacral"],
    }
    
    for key, keywords in keyword_map.items():
        for kw in keywords:
            if kw in text_lower:
                analysis[key] = True
                break
    
    return analysis

# ============================================================
# DOCX OKUMA
# ============================================================
def read_docx(filepath):
    """python-docx ile DOCX oku"""
    try:
        from docx import Document
        doc = Document(filepath)
        text_parts = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text_parts.append(para.text)
        return len(doc.paragraphs), "\n".join(text_parts)
    except Exception as e:
        return -1, f"DOCX HATA: {str(e)}"

# ============================================================
# RESIM OKUMA
# ============================================================
def read_image(filepath):
    """Pillow ile resim oku, meta veri çıkar"""
    try:
        from PIL import Image
        img = Image.open(filepath)
        info = {
            "format": img.format,
            "mode": img.mode,
            "size": img.size,
            "width": img.width,
            "height": img.height,
            "dpi": img.info.get("dpi", None),
            "exif": {}
        }
        # EXIF verisi
        try:
            exif_data = img._getexif()
            if exif_data:
                for tag_id, value in exif_data.items():
                    info["exif"][str(tag_id)] = str(value)[:200]
        except:
            pass
        return info
    except Exception as e:
        return {"error": str(e)}

# ============================================================
# ANA İŞLEM
# ============================================================
def main():
    print("=" * 70)
    print("FAZ-3: PDF/DOCX/RESIM DOSYALARI OKUYUCU")
    print("=" * 70)
    
    # PDF'leri tara
    pdf_files = [f for f in os.listdir(REPO_DIR) if f.lower().endswith('.pdf')]
    results["summary"]["total_pdfs"] = len(pdf_files)
    print(f"\n>>> PDF Dosyaları taranıyor ({len(pdf_files)} adet)...")
    
    for pdf_file in sorted(pdf_files):
        filepath = os.path.join(REPO_DIR, pdf_file)
        print(f"  [{pdf_file[:60]}...] ", end="", flush=True)
        
        # Önce PyPDF2 ile dene
        num_pages, text = read_pdf_pypdf2(filepath)
        method = "PyPDF2"
        
        # PyPDF2 başarısız olursa pdfplumber dene
        if num_pages < 0 or len(text) < 50:
            num_pages2, text2 = read_pdf_pdfplumber(filepath)
            if num_pages2 > 0 and len(text2) > len(text):
                num_pages, text = num_pages2, text2
                method = "pdfplumber"
        
        if num_pages > 0 and len(text) > 10:
            numerical = extract_numerical_values(text)
            analysis = analyze_pdf_content(text)
            results["pdf_files"][pdf_file] = {
                "status": "OK",
                "method": method,
                "num_pages": num_pages,
                "text_length": len(text),
                "text_preview": text[:3000],
                "numerical_findings": numerical,
                "analysis": {k: v for k, v in analysis.items() if v}
            }
            results["summary"]["pdfs_read"] += 1
            print(f"OK ({num_pages} sayfa, {len(text)} karakter, {method})")
        else:
            results["pdf_files"][pdf_file] = {
                "status": "FAILED",
                "method": method,
                "num_pages": num_pages,
                "error": text[:500]
            }
            results["summary"]["pdfs_failed"] += 1
            print(f"HATA: {text[:80]}")
    
    # DOCX'leri tara
    docx_files = [f for f in os.listdir(REPO_DIR) if f.lower().endswith('.docx')]
    results["summary"]["total_docxs"] = len(docx_files)
    print(f"\n>>> DOCX Dosyaları taranıyor ({len(docx_files)} adet)...")
    
    for docx_file in sorted(docx_files):
        filepath = os.path.join(REPO_DIR, docx_file)
        print(f"  [{docx_file[:60]}...] ", end="", flush=True)
        
        num_paras, text = read_docx(filepath)
        
        if num_paras > 0 and len(text) > 10:
            numerical = extract_numerical_values(text)
            analysis = analyze_pdf_content(text)
            results["docx_files"][docx_file] = {
                "status": "OK",
                "num_paragraphs": num_paras,
                "text_length": len(text),
                "text_preview": text[:3000],
                "numerical_findings": numerical,
                "analysis": {k: v for k, v in analysis.items() if v}
            }
            results["summary"]["docxs_read"] += 1
            print(f"OK ({num_paras} paragraf, {len(text)} karakter)")
        else:
            results["docx_files"][docx_file] = {
                "status": "FAILED",
                "num_paragraphs": num_paras,
                "error": text[:500]
            }
            results["summary"]["docxs_failed"] += 1
            print(f"HATA: {text[:80]}")
    
    # Resimleri tara
    img_extensions = ('.jpg', '.jpeg', '.png', '.jfif', '.gif', '.bmp', '.webp')
    img_files = [f for f in os.listdir(REPO_DIR) if f.lower().endswith(img_extensions)]
    results["summary"]["total_images"] = len(img_files)
    print(f"\n>>> Resim Dosyaları taranıyor ({len(img_files)} adet)...")
    
    for img_file in sorted(img_files):
        filepath = os.path.join(REPO_DIR, img_file)
        print(f"  [{img_file[:60]}...] ", end="", flush=True)
        
        info = read_image(filepath)
        
        if "error" not in info:
            results["image_files"][img_file] = {
                "status": "OK",
                **info
            }
            results["summary"]["images_read"] += 1
            print(f"OK ({info['format']} {info['width']}x{info['height']})")
        else:
            results["image_files"][img_file] = {
                "status": "FAILED",
                "error": info["error"]
            }
            results["summary"]["images_failed"] += 1
            print(f"HATA: {info['error'][:80]}")
    
    # JSON'a kaydet
    print(f"\n>>> Sonuçlar JSON'a kaydediliyor: {OUTPUT_JSON}")
    with open(OUTPUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    # Özet
    s = results["summary"]
    print("\n" + "=" * 70)
    print("ÖZET")
    print("=" * 70)
    print(f"  PDF:  {s['pdfs_read']}/{s['total_pdfs']} okundu, {s['pdfs_failed']} başarısız")
    print(f"  DOCX: {s['docxs_read']}/{s['total_docxs']} okundu, {s['docxs_failed']} başarısız")
    print(f"  IMG:  {s['images_read']}/{s['total_images']} okundu, {s['images_failed']} başarısız")
    print(f"  JSON: {OUTPUT_JSON}")
    print("=" * 70)
    
    return results

if __name__ == "__main__":
    main()
