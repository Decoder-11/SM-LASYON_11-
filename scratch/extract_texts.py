#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Extract text from PDF and DOCX files and save as UTF-8 text files.
Handles Turkish characters properly.
"""

import subprocess
import sys
import os

def install_if_missing(package, import_name=None):
    """Install a package if it's not already installed."""
    if import_name is None:
        import_name = package
    try:
        __import__(import_name)
        print(f"  ✓ {package} already installed")
    except ImportError:
        print(f"  Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package],
                              stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        print(f"  ✓ {package} installed successfully")

# Ensure required packages are installed
print("Checking dependencies...")
install_if_missing("pdfplumber")
install_if_missing("python-docx", "docx")

import pdfplumber
from docx import Document

# Define the scratch output directory
SCRATCH_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(SCRATCH_DIR, exist_ok=True)

def extract_pdf(pdf_path, output_path):
    """Extract all text from a PDF using pdfplumber, save as UTF-8 text."""
    print(f"\n{'='*60}")
    print(f"Processing PDF: {os.path.basename(pdf_path)}")
    print(f"{'='*60}")
    
    if not os.path.exists(pdf_path):
        print(f"  ✗ ERROR: File not found: {pdf_path}")
        return False
    
    try:
        all_text = []
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            print(f"  Total pages: {total_pages}")
            
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    all_text.append(f"--- Sayfa {i+1} / Page {i+1} ---\n")
                    all_text.append(page_text)
                    all_text.append("\n\n")
                    print(f"  Page {i+1}/{total_pages}: {len(page_text)} chars extracted")
                else:
                    all_text.append(f"--- Sayfa {i+1} / Page {i+1} ---\n")
                    all_text.append("[Bu sayfada metin bulunamadı / No text found on this page]\n\n")
                    print(f"  Page {i+1}/{total_pages}: No text found (might be image-based)")
        
        full_text = "".join(all_text)
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        
        print(f"  ✓ Saved to: {output_path}")
        print(f"  Total characters: {len(full_text)}")
        return True
        
    except Exception as e:
        print(f"  ✗ ERROR extracting PDF: {e}")
        return False

def extract_docx(docx_path, output_path):
    """Extract all text from a DOCX using python-docx, save as UTF-8 text."""
    print(f"\n{'='*60}")
    print(f"Processing DOCX: {os.path.basename(docx_path)}")
    print(f"{'='*60}")
    
    if not os.path.exists(docx_path):
        print(f"  ✗ ERROR: File not found: {docx_path}")
        return False
    
    try:
        doc = Document(docx_path)
        all_text = []
        
        # Extract text from paragraphs
        para_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
                para_count += 1
        
        # Also extract text from tables
        table_count = 0
        for table in doc.tables:
            table_count += 1
            all_text.append(f"\n--- Tablo {table_count} / Table {table_count} ---")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    all_text.append(" | ".join(row_text))
        
        full_text = "\n".join(all_text)
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        
        print(f"  ✓ Saved to: {output_path}")
        print(f"  Paragraphs: {para_count}, Tables: {table_count}")
        print(f"  Total characters: {len(full_text)}")
        return True
        
    except Exception as e:
        print(f"  ✗ ERROR extracting DOCX: {e}")
        return False

# ============================================================
# Main execution
# ============================================================
print("\n" + "="*60)
print("PDF & DOCX Text Extraction Script")
print("="*60)

# File 1: BİO REZONAS-FREKANS PDF
pdf1_path = r"C:\Users\soldi\IdeaProjects\simülation-11\BİO REZONAS-FREKANS PDF..pdf"
pdf1_output = os.path.join(SCRATCH_DIR, "bio_rezonans_extracted.txt")
result1 = extract_pdf(pdf1_path, pdf1_output)

# File 2: HÜD HÜD VE FREKANS PDF
pdf2_path = r"C:\Users\soldi\IdeaProjects\simülation-11\HÜD HÜD VE FREKANS PDF..pdf"
pdf2_output = os.path.join(SCRATCH_DIR, "hudhud_frekans_extracted.txt")
result2 = extract_pdf(pdf2_path, pdf2_output)

# File 3: gemini-levhi mahfuz.docx (found on Desktop)
docx_path = r"C:\Users\soldi\OneDrive\Masaüstü\gemini-levhi mahfuz.docx"
docx_output = os.path.join(SCRATCH_DIR, "gemini_levhi_mahfuz_extracted.txt")
result3 = extract_docx(docx_path, docx_output)

# Summary
print("\n" + "="*60)
print("SUMMARY")
print("="*60)
print(f"  1. BİO REZONAS-FREKANS:    {'✓ SUCCESS' if result1 else '✗ FAILED'}")
print(f"  2. HÜD HÜD VE FREKANS:     {'✓ SUCCESS' if result2 else '✗ FAILED'}")
print(f"  3. gemini-levhi mahfuz:     {'✓ SUCCESS' if result3 else '✗ FAILED'}")
print(f"\nOutput directory: {SCRATCH_DIR}")
print("="*60)
